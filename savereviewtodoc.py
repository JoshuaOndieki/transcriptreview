import imp
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH


def savereview(file_id, transcript):
    document = Document("templates/review_template.docx")

    document.add_heading('REVIEW: ' + file_id, 0)

    document.add_paragraph('Transcript review')
    document.add_paragraph().add_run('Yellow is for extras not discussed in the basic guidelines. Understand their usage if you are interested in QA work. Your metrics does not include them.').italic = True
    document.add_paragraph('Please note as at now I have not included speaker corrections. I have decided to create reviews so that you can note the reoccurring errors and hopefully I will get a reviewer or two.')

    removal_p = document.add_paragraph()
    removal_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    addition_p = document.add_paragraph()
    addition_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rating_p = document.add_paragraph()
    rating_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # sum (1 for x in transcript if x[0] != "?")

    word_count = 0
    removal = 0
    addition = 0

    for d in transcript:
        if d[0] == "?":
            continue

        if d[-1] == ":":
           p = document.add_paragraph()
           d = d[1:]
           p.add_run(d + " ").bold = True

        elif ":" in d and d[-1] != ":":
            d = d[1:]
            p.add_run(d + " ").bold = True
        else:
            if d[0] == "-":
                d = d[1:]
                font = p.add_run(d).font
                if "<" and ">" in d:
                    font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    word_count += 1
                    removal += 1
                    font.highlight_color = WD_COLOR_INDEX.RED
            elif d[0] == "+":
                d = d[1:]
                font = p.add_run(d).font
                if "<" and ">" in d:
                    font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    word_count += 1
                    addition += 1
                    font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            else:
                word_count += 1
                p.add_run(d)

    removal_font = removal_p.add_run(" REMOVAL: " + str(removal)).font
    removal_font.bold = True
    removal_font.color.rgb = RGBColor(255, 0, 0)
    

    addition_font = addition_p.add_run("  ADDITION: " + str(addition)).font
    addition_font.bold = True
    addition_font.color.rgb = RGBColor(0, 122, 55)

    rating_font = rating_p.add_run("QUALITY:   " + str(100 - round(((addition / word_count * 100) + (removal / word_count * 100)) / 2.2, 2)) + "%").font
    rating_font.bold = True

    document.save("files/docx/" + file_id + "-review.docx")